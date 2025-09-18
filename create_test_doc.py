from docx import Document

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('فایل تستی Word', 0)

# Add a paragraph
doc.add_paragraph('این یک فایل Word تستی است که برای تست قابلیت پردازش فایل‌های رایج در سیستم MobixAI ایجاد شده است.')

# Save the document
doc.save('test_file.docx')